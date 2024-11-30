from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework import status
from django.http import HttpResponse
from docx import Document
from .models import Resume
from .serializers import ResumeSerializer
from django.utils.html import strip_tags
from django.contrib.auth.models import User
from rest_framework_simplejwt.tokens import RefreshToken
from .models import UserProfile
from xhtml2pdf import pisa
from django.template.loader import render_to_string
from docx import Document
from docx.shared import Pt
from io import BytesIO
from bs4 import BeautifulSoup
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class UserRegistrationView(APIView):
    authentication_classes = []
    permission_classes = []

    def post(self, request):
        username = request.data.get("username")
        password = request.data.get("password")
        name = request.data.get("name")
        email = request.data.get("email")
        age = request.data.get("age")
        dob = request.data.get("dob")
        gender = request.data.get("gender")

        if User.objects.filter(username=username).exists():
            return Response({"error": "Username already exists."}, status=status.HTTP_400_BAD_REQUEST)

        user = User.objects.create_user(username=username, password=password, email=email)
        UserProfile.objects.create(
            user=user,
            username=username,
            name = name,
            email = email,
            age=age,
            dob=dob,
            gender=gender,
        )
        default_resume = Resume.objects.create(
            user=user,
            title="Default Resume",
            resume_content="Enter your Resume content here...",
        )
        default_resume.save()

        return Response({"message": "User registered successfully"}, status=status.HTTP_201_CREATED)

class UserProfileView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        print(request)
        profile = request.user.profile
        data = {
            "username": request.user.username,
            "email": request.user.email,
            "age": profile.age,
            "dob": profile.dob,
            "gender": profile.gender,
            "name":profile.name
        }
        return Response(data)

    def put(self, request):
        profile = request.user.profile
        profile.age = request.data.get("age", profile.age)
        profile.dob = request.data.get("dob", profile.dob)
        profile.gender = request.data.get("gender", profile.gender)
        profile.name = request.data.get("name", profile.name)
        profile.save()
        return Response({"message": "Profile updated successfully"})


class ResumeView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, pk=None):
        if pk:
            try:
                resume = Resume.objects.get(pk=pk, user=request.user)
                serializer = ResumeSerializer(resume)
                return Response(serializer.data)
            except Resume.DoesNotExist:
                return Response({'error': 'Resume not found'}, status=status.HTTP_404_NOT_FOUND)

        resumes = Resume.objects.filter(user=request.user)
        serializer = ResumeSerializer(resumes, many=True)
        return Response(serializer.data)

    def post(self, request):
        data = request.data
        data['user'] = request.user.id
        serializer = ResumeSerializer(data=data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def put(self, request, pk):
        try:
            resume = Resume.objects.get(pk=pk, user=request.user)
        except Resume.DoesNotExist:
            return Response({'error': 'Resume not found'}, status=status.HTTP_404_NOT_FOUND)

        serializer = ResumeSerializer(resume, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk):
        try:
            resume = Resume.objects.get(pk=pk, user=request.user)
        except Resume.DoesNotExist:
            return Response({'error': 'Resume not found'}, status=status.HTTP_404_NOT_FOUND)

        resume.delete()
        return Response({'message': 'Resume deleted successfully'}, status=status.HTTP_204_NO_CONTENT)

class GeneratePDF(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, pk):
        try:
            resume = Resume.objects.get(pk=pk, user=request.user)
        except Resume.DoesNotExist:
            return HttpResponse("Resume not found", status=404)

        resume_content = resume.resume_content 

        context = {
            "resume_content": resume_content,  
        }
        html = render_to_string("resume_content_template.html", context)
        
        response = HttpResponse(content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{resume.title}.pdf"'

        pisa_status = pisa.CreatePDF(html, dest=response)
        if pisa_status.err:
            return HttpResponse("Error generating PDF", status=500)
        return response


class GenerateDOCX(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, pk):
        try:
            resume = Resume.objects.get(pk=pk, user=request.user)
        except Resume.DoesNotExist:
            return HttpResponse("Resume not found", status=404)
        context = {
            "resume_content": resume.resume_content,
        }
        rendered_html = render_to_string("resume_content_template.html", context)

        soup = BeautifulSoup(rendered_html, "html.parser")

        document = Document()

        for element in soup.body.children:
            if element.name == "h1":
                paragraph = document.add_heading(element.get_text(strip=True), level=1)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif element.name == "h2":
                document.add_heading(element.get_text(strip=True), level=2)
            elif element.name == "h3":
                document.add_heading(element.get_text(strip=True), level=3)
            elif element.name == "p":
                document.add_paragraph(element.get_text(strip=True))
            elif element.name == "ul":
                for li in element.find_all("li"):
                    document.add_paragraph(f"• {li.get_text(strip=True)}", style="List Bullet")
            elif element.name == "ol":
                for li in element.find_all("li"):
                    document.add_paragraph(f"{li.get_text(strip=True)}", style="List Number")
            elif element.name == "strong":
                paragraph = document.add_paragraph()
                run = paragraph.add_run(element.get_text(strip=True))
                run.bold = True
            elif element.name == "em":
                paragraph = document.add_paragraph()
                run = paragraph.add_run(element.get_text(strip=True))
                run.italic = True

        response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response["Content-Disposition"] = f'attachment; filename="{resume.title}.docx"'
        docx_stream = BytesIO()
        document.save(docx_stream)
        docx_stream.seek(0)
        response.write(docx_stream.read())

        return response


class LogoutView(APIView):
    def post(self, request):
        try:
            refresh_token = request.data["refresh"]
            token = RefreshToken(refresh_token)
            print(token)
            token.blacklist()
            print("after blacklist")
            return Response({"message": "Logged out successfully"}, status=200)
        except Exception as e:
            return Response({"error": str(e)}, status=400)

